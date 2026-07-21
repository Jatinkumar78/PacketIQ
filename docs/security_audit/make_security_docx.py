#!/usr/bin/env python3
"""Render the PacketIQ security audit into a professional Word (.docx) report.

Reuses the curated findings/controls/steps data from make_security_pdf.py so the
Word and PDF versions never drift.
"""
from __future__ import annotations

import os
import re
import sys
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import make_security_pdf as data  # noqa: E402

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "reports", "PacketIQ_Security_Audit_Report.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x47, 0x55, 0x69)
SEV_HEX = {"CRITICAL": "B91C1C", "HIGH": "C2410C", "MEDIUM": "A16207", "LOW": "15803D", "INFO": "475569"}
NAVY_HEX = "0B1F3A"
LIGHT_HEX = "EEF2F7"
GREEN_HEX = "15803D"


def _strip(html):
    """Remove the light HTML markup used in the PDF data (<font>, <b>, &nbsp;, etc.)."""
    t = re.sub(r"<[^>]+>", "", html)
    return (t.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<")
            .replace("&gt;", ">"))


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_fill})
    tcPr.append(sh)


def set_cell(cell, text, *, bold=False, color=None, size=9, fill=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if fill:
        shade(cell, fill)
    for pr in cell.paragraphs:
        pr.paragraph_format.space_after = Pt(1)
        pr.paragraph_format.space_before = Pt(1)


def hrule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6", qn("w:space"): "1", qn("w:color"): "1D4ED8"})
    pbdr.append(bottom)
    pPr.append(pbdr)


def build():
    doc = Document()
    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    for h, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        st = doc.styles[h]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.color.rgb = NAVY if h == "Heading 1" else BLUE

    counts = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0, "INFO": 0}
    for f in data.FINDINGS:
        counts[f["sev"]] += 1

    # ── Title block ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PacketIQ")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Security Audit Report")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = BLUE
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("AI PCAP Forensics & SOC Copilot — Static, Dependency & Secret Assessment")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    doc.add_paragraph()

    # severity chips as a 1x5 table
    chip = doc.add_table(rows=2, cols=5)
    chip.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, sev in enumerate(("CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO")):
        set_cell(chip.cell(0, i), str(counts[sev]), bold=True, color=WHITE, size=16,
                 fill=SEV_HEX[sev], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(chip.cell(1, i), sev, bold=True, color=WHITE, size=8,
                 fill=SEV_HEX[sev], align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    meta = [
        ("Target", "PacketIQ v1.0.0 — Python network-forensics tool (CLI + FastAPI web app)"),
        ("Audit type", "White-box: manual code review + SAST (bandit) + dependency CVE scan (pip-audit) + git secret scan + live exploitation"),
        ("Scope", "All Python source under packetiq/, dependencies, git history. Authorised self-assessment of the author's own project."),
        ("Environment", f"Sandboxed local venv · Python {__import__('platform').python_version()} · {__import__('platform').system()}"),
        ("Date", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
        ("Result", f"{counts['CRITICAL']} Critical, {counts['MEDIUM']} Medium, {counts['LOW']} Low, "
                   f"{counts['INFO']} Info. All findings remediated: every code-level fix is in place and the "
                   "Critical (leaked keys) was closed by owner key-revocation on 2026-07-12."),
    ]
    mt = doc.add_table(rows=len(meta), cols=2)
    mt.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        set_cell(mt.cell(i, 0), k, bold=True, size=9, fill=LIGHT_HEX)
        set_cell(mt.cell(i, 1), v, size=9)
    mt.columns[0].width = Pt(90)

    doc.add_page_break()

    # ── 1. Executive summary ──
    doc.add_heading("1. Executive Summary", level=1)
    hrule(doc)
    doc.add_paragraph(_strip(
        "PacketIQ was assessed with a white-box methodology combining manual source review of every "
        "security-sensitive code path with automated tooling (bandit SAST, pip-audit dependency scan, "
        "git-history secret scan) and — in a second round — live exploitation attempts against the running "
        "server. The codebase is fundamentally sound: parameterised SQL, file paths validated against a "
        "server-side job registry, no dangerous execution sinks, timeouts on every outbound request, and "
        "HTML-escaping of all rendered data."))
    doc.add_paragraph(_strip(
        "Twelve findings were raised across two rounds. The single Critical was operational, not code: live AI "
        "keys committed to git history. The account owner has since revoked those keys (2026-07-12), so they are "
        "now dead and the finding is closed, with an optional history scrub remaining only for tidiness. Six Medium "
        "findings were fixed — an upload memory-exhaustion DoS, oversized HTTP buffers, vulnerable dependencies, and "
        "(found by actively attempting the exploits) a path-traversal / arbitrary file write, DNS-rebinding, and CSRF "
        "against the local server. Three Low findings were also fixed. Every code-level fix is covered by a new "
        "regression test and, where exploitable, re-tested live to confirm it is blocked."))

    doc.add_heading("Findings overview", level=2)
    ov = doc.add_table(rows=1, cols=4)
    ov.style = "Table Grid"
    for i, h in enumerate(("ID", "Finding", "Severity", "Status")):
        set_cell(ov.cell(0, i), h, bold=True, color=WHITE, size=9, fill=NAVY_HEX)
    for f in data.FINDINGS:
        row = ov.add_row().cells
        set_cell(row[0], f["id"], size=9)
        set_cell(row[1], f["title"], size=9)
        set_cell(row[2], f["sev"], bold=True, size=9, color=RGBColor.from_string(SEV_HEX[f["sev"]]))
        set_cell(row[3], _strip(f["status"]).split("—")[0].split("(")[0].strip(), size=8)
    for w, c in zip((50, 230, 60, 130), ov.columns):
        c.width = Pt(w)

    doc.add_page_break()

    # ── 2. Methodology ──
    doc.add_heading("2. Scope & Methodology", level=1)
    hrule(doc)
    doc.add_paragraph(_strip(
        "This is an authorised, defensive self-assessment of the author's own project, performed entirely in a "
        "local sandbox. No external systems, networks or third parties were touched. Four complementary "
        "techniques were used: manual review of the highest-risk paths; bandit static analysis; pip-audit "
        "dependency CVE scanning; a git-history secret scan; and live exploitation attempts to confirm each "
        "class rather than assume it safe."))

    # ── 3. Procedure ──
    doc.add_heading("3. Audit Procedure — Step by Step", level=1)
    hrule(doc)
    doc.add_paragraph("Each step records why it was performed, the command/action, and the result.").runs[0].font.size = Pt(9)
    stt = doc.add_table(rows=1, cols=4)
    stt.style = "Table Grid"
    for i, h in enumerate(("#", "Step & why", "Command / action", "Result")):
        set_cell(stt.cell(0, i), h, bold=True, color=WHITE, size=9, fill=NAVY_HEX)
    for i, (what, cmd, res, why) in enumerate(data.STEPS, 1):
        row = stt.add_row().cells
        set_cell(row[0], str(i), size=9)
        cell = row[1]
        cell.text = ""
        p = cell.paragraphs[0]
        rr = p.add_run(what)
        rr.bold = True
        rr.font.size = Pt(9)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(why)
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = GREY
        set_cell(row[2], cmd, size=8)
        set_cell(row[3], res, size=8)
    for w, c in zip((26, 170, 150, 130), stt.columns):
        c.width = Pt(w)

    doc.add_page_break()

    # ── 4. Findings detail ──
    doc.add_heading("4. Findings in Detail", level=1)
    hrule(doc)
    for f in data.FINDINGS:
        doc.add_heading(f"{f['id']} · {f['title']}", level=2)
        dt = doc.add_table(rows=3, cols=2)
        dt.style = "Table Grid"
        set_cell(dt.cell(0, 0), "Severity", bold=True, size=9, fill=LIGHT_HEX)
        set_cell(dt.cell(0, 1), f"{f['sev']} · {f['cwe']}", bold=True, size=9,
                 color=RGBColor.from_string(SEV_HEX[f["sev"]]))
        set_cell(dt.cell(1, 0), "Description", bold=True, size=9, fill=LIGHT_HEX)
        set_cell(dt.cell(1, 1), _strip(f["desc"]), size=9)
        set_cell(dt.cell(2, 0), "Impact", bold=True, size=9, fill=LIGHT_HEX)
        set_cell(dt.cell(2, 1), _strip(f["impact"]), size=9)
        dt.columns[0].width = Pt(70)
        ep = doc.add_paragraph()
        er = ep.add_run("Evidence:  ")
        er.bold = True
        er.font.size = Pt(8.5)
        ev = ep.add_run(_strip(f["evidence"]))
        ev.font.name = "Consolas"
        ev.font.size = Pt(8)
        ft = doc.add_table(rows=2, cols=2)
        ft.style = "Table Grid"
        set_cell(ft.cell(0, 0), "Remediation", bold=True, size=9, fill=LIGHT_HEX)
        set_cell(ft.cell(0, 1), _strip(f["fix"]), size=9)
        set_cell(ft.cell(1, 0), "Status", bold=True, size=9, fill=LIGHT_HEX)
        set_cell(ft.cell(1, 1), _strip(f["status"]), size=9, color=RGBColor.from_string(GREEN_HEX))
        ft.columns[0].width = Pt(70)
        doc.add_paragraph()

    doc.add_page_break()

    # ── 5. Controls verified ──
    doc.add_heading("5. Security Controls Verified (No Finding)", level=1)
    hrule(doc)
    ct = doc.add_table(rows=1, cols=2)
    ct.style = "Table Grid"
    set_cell(ct.cell(0, 0), "Control", bold=True, color=WHITE, size=9, fill=GREEN_HEX)
    set_cell(ct.cell(0, 1), "Result", bold=True, color=WHITE, size=9, fill=GREEN_HEX)
    for k, v in data.CONTROLS_OK:
        row = ct.add_row().cells
        set_cell(row[0], k, bold=True, size=9)
        set_cell(row[1], "✓ " + _strip(v), size=8.5)
    ct.columns[0].width = Pt(130)

    # ── 6. Verification + required action ──
    doc.add_page_break()
    doc.add_heading("6. Remediation Verification & Required Owner Action", level=1)
    hrule(doc)
    doc.add_paragraph("After applying the fixes the full battery was re-run:").runs[0].font.size = Pt(10)
    ver = [
        ("Unit/integration tests", "304 passed — the full suite, including the security regression tests added in this engagement — no regressions."),
        ("Linter (ruff)", "All checks passed."),
        ("bandit (SAST)", "High: 0, Medium: 0, Low: 53. Remaining Low findings are defensive try/except/pass and list-form subprocess calls (no shell=True). The pseudo-random calls in the synthetic capture/benchmark generators are annotated \"# nosec B311\" (deterministic sample data, never cryptographic)."),
        ("Live exploit re-tests", "Path traversal (ip=../../..) → 400, no file written; DNS-rebinding (bad Host) → 400; "
                                  "cross-origin POST to privileged setup-capture → 403."),
        ("pip-audit", "Reference environment migrated to Python 3.12.13: the pinned floors resolve to fully-patched "
                      "releases and the runtime dependency set reports zero advisories. One dev-only transitive remains "
                      "(diskcache via the dev-extra pySigma), no upstream fix yet. Python 3.9 stays supported at newest-compatible pins."),
    ]
    vt = doc.add_table(rows=len(ver), cols=2)
    vt.style = "Table Grid"
    for i, (k, v) in enumerate(ver):
        set_cell(vt.cell(i, 0), k, bold=True, size=9, fill=LIGHT_HEX)
        set_cell(vt.cell(i, 1), v, size=9)
    vt.columns[0].width = Pt(120)

    doc.add_heading("Owner action (F-01) — completed", level=2)
    p = doc.add_paragraph()
    p.add_run("The account owner revoked both exposed keys in the provider consoles (Google AI Studio & Groq) on "
              "2026-07-12, so the leaked values no longer authenticate. Optionally, scrub them from git history before "
              "any public push so the dead values do not linger in the record:").font.size = Pt(10)
    code = doc.add_paragraph()
    cr = code.add_run("pip install git-filter-repo\n"
                      "git filter-repo --path .env.example --invert-paths\n"
                      "git push --force --all")
    cr.font.name = "Consolas"
    cr.font.size = Pt(9)

    # ── 7. Conclusion ──
    doc.add_heading("7. Conclusion", level=1)
    hrule(doc)
    doc.add_paragraph(_strip(
        "At the code level PacketIQ is a well-built and defensible application: it avoids the common high-impact "
        "web vulnerabilities and now streams uploads, bounds its buffers, validates Host/Origin headers, blocks "
        "path traversal, hardens its one privileged path, and restricts its data directories. After this "
        "engagement there are no outstanding code-level vulnerabilities, including a real path-traversal and "
        "CSRF/DNS-rebinding class closed in round two. The one operational risk, API keys committed to git "
        "history, has been closed by the owner revoking those keys. The reference/dev environment has been migrated to "
        "Python 3.12.13, so the pinned dependency floors now resolve to fully-patched releases (Python 3.9 stays supported "
        "for anyone who needs it); further hardening (authentication, rate-limiting) is worthwhile only if the web app is ever exposed beyond localhost."))

    # footer
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    fr = footer.add_run("PacketIQ Security Audit Report · Confidential · Authorised self-assessment of the author's own project")
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = GREY

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
